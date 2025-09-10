import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import pandas as pd
import json
import os
import re
import requests
import argparse
from dotenv import load_dotenv
from unstructured.partition.xlsx import partition_xlsx
from docx import Document  # New for Word

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"

# Configuring Tesseract path (change as needed for your OS)
pytesseract.pytesseract.tesseract_cmd = r'C:\Users\Shreyas\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'

# === OCR Extractors ===
def extract_text_from_image(image_path: str) -> str:
    image = Image.open(image_path)
    return pytesseract.image_to_string(image).strip()

def extract_text_from_pdf(pdf_path: str) -> str:
    pages = convert_from_path(pdf_path, dpi=300)
    print(f"[📄] Extracting text from {len(pages)} PDF pages...")
    full_text = ""
    for i, page in enumerate(pages):
        text = pytesseract.image_to_string(page)
        full_text += f"\n\n--- Page {i + 1} ---\n{text}"
    return full_text.strip()

# def extract_text_from_excel(excel_path: str, structure_threshold: float = 0.5) -> str:
#     try:
#         # dfs = pd.read_excel(excel_path, sheet_name=None, engine='openpyxl')
#         # total_cells = 0
#         # empty_cells = 0

#         # for sheet, df in dfs.items():
#         #     total_cells += df.size
#         #     empty_cells += df.isna().sum().sum()

#         # sparsity = empty_cells / total_cells if total_cells > 0 else 1

#         # if sparsity < structure_threshold:
#         #     full_text = ""
#         #     for name, df in dfs.items():
#         #         full_text += f"\n\n--- Sheet: {name} ---\n"
#         #         full_text += df.fillna("").to_string(index=False)
#         #     print("📊 Structured Excel detected. Parsed using pandas.")
#         #     return full_text.strip()
#         if True:
#             print("📎 Unstructured Excel detected. Parsing using `unstructured`...")
#             elements = partition_xlsx(filename=excel_path)
#             return "\n".join([e.text for e in elements if e.text]).strip()

    # except Exception as e:
    #     print(f"❌ Excel processing failed: {e}")
    #     return ""

# def extract_text_from_word(word_path: str) -> str:
#     try:
#         doc = Document(word_path)
#         print(f"[📄] Extracting text from Word document: {word_path}")
#         text_parts = []

#         for table in doc.tables:
#             rows = []
#             for row in table.rows:
#                 cells = [cell.text.strip().replace("\n"," ") for cell in row.cells]
#                 rows.append("\t".join(cells))  # use tab as separator
#             text_parts.append("\n".join(rows))

#             # Optionally: Extract non-table paragraphs
#         paragraph_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#          # Combine table text + any paragraph text (order based on your needs)
#         return "\n\n".join(text_parts + [paragraph_text]).strip()
#     except Exception as e:
#         print(f"❌ Word document parsing failed: {e}")
#         return ""

# === Chunking ===
def chunk_text(text: str, max_tokens: int = 800) -> list:
    lines = text.splitlines()
    chunks, chunk = [], ""

    for line in lines:
        if len((chunk + line).split()) > max_tokens:
            chunks.append(chunk.strip())
            chunk = line + "\n"
        else:
            chunk += line + "\n"

    if chunk:
        chunks.append(chunk.strip())

    return chunks

# === Prompt Builder ===
def create_prompt(text: str, schema: dict = None) -> str:
    if schema:
        return f"""
You are an intelligent document parser.

First, carefully analyze the following text step-by-step and identify important fields.

Then, extract structured data in valid JSON format using the schema below.

Schema:
{json.dumps(schema, indent=2)}

Text:
\"\"\"{text}\"\"\"

Now return only the valid JSON output matching the schema. No explanation or reasoning steps.
"""
    else:
        return f"""
You are a JSON extraction agent. Convert the following text into structured JSON.

Text:
\"\"\"{text}\"\"\"

Only return valid JSON. No explanation.
"""

# === Groq LLM Call ===
def call_groq_llm(prompt: str, model: str = "llama3-70b-8192"):
    if not GROQ_API_KEY:
        raise EnvironmentError("❌ GROQ_API_KEY is not set.")

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2
    }

    response = requests.post(GROQ_URL, headers=headers, json=payload)
    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        raise Exception(f"❌ Groq API error {response.status_code}: {response.text}")

# === JSON Cleaner ===
def extract_json_from_llm_output(raw_output: str):
    try:
        json_match = re.search(r"{[\s\S]+}", raw_output)
        if json_match:
            return json.loads(json_match.group(0))
        else:
            raise ValueError("No valid JSON found.")
    except Exception as e:
        print("[❌] Failed to parse JSON:", e)
        return {"raw_output": raw_output}

# === Main Pipeline ===
def run_extraction(path: str, schema=None, model="llama3-70b-8192"):
    print(f"🔍 Extracting from: {path}")
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        text = extract_text_from_image(path)
    # elif ext in [".xls", ".xlsx"]:
    #     text = extract_text_from_excel(path)
    # elif ext in [".doc", ".docx"]:
    #     text = extract_text_from_word(path)
    else:
        raise ValueError("Unsupported file format. Only PDF, images, Excel, and Word supported.")

    print("✂️ Chunking extracted text...")
    chunks = chunk_text(text)

    print(f"🧩 Total chunks: {len(chunks)}")
    results = []

    for i, chunk in enumerate(chunks):
        print(f"💬 Sending chunk {i+1} to Groq LLM...")
        prompt = create_prompt(chunk, schema)
        try: 
            raw_output = call_groq_llm(prompt, model)
            cleaned = extract_json_from_llm_output(raw_output)
            results.append(cleaned)
        except Exception as e:
            print(f"❌ Error while processing chunk {i+1}: {e}")
            print("⚠️ Stopping further processing due to API limit or error.")
            break

    final_result = {
        "file": os.path.basename(path),
        "chunks": len(chunks),
        "results": results
    }

    print("✅ Final Structured Output:")
    print(json.dumps(final_result, indent=2))
    return final_result

# === CLI Entry ===
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("path", help="Path to input file (image, PDF, Excel, or Word)")
    parser.add_argument("--schema", help="Optional path to schema JSON file")
    parser.add_argument("--model", default="llama3-70b-8192", help="Groq model to use")
    args = parser.parse_args()

    schema = None
    if args.schema and os.path.exists(args.schema):
        with open(args.schema, "r") as f:
            schema = json.load(f)

    run_extraction(args.path, schema=schema, model=args.model)
